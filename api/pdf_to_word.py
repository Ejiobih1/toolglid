"""
Advanced PDF to Word Converter
High-quality conversion preserving text, images, tables, and formatting.

Uses:
- PyMuPDF (fitz) for accurate PDF parsing
- python-docx for Word document creation
- Custom algorithms for structure detection

Author: ToolGlid
"""

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import io
import re
from dataclasses import dataclass
from collections import defaultdict


@dataclass
class TextBlock:
    """Represents a block of text with formatting"""
    text: str
    x0: float
    y0: float
    x1: float
    y1: float
    font_name: str
    font_size: float
    color: Tuple[float, float, float]
    is_bold: bool
    is_italic: bool
    block_no: int
    line_no: int


@dataclass
class ImageBlock:
    """Represents an image in the PDF"""
    image_data: bytes
    x0: float
    y0: float
    x1: float
    y1: float
    width: float
    height: float
    ext: str


@dataclass
class TableCell:
    """Represents a table cell"""
    text: str
    row: int
    col: int
    x0: float
    y0: float
    x1: float
    y1: float
    bg_color: Tuple[float, float, float] = None  # Background color (RGB 0-1)
    text_color: Tuple[float, float, float] = None  # Text color (RGB 0-1)
    font_size: float = 11
    is_bold: bool = False


class PDFToWordConverter:
    """
    High-quality PDF to Word converter.

    Features:
    - Text extraction with font, size, color preservation
    - Image extraction and embedding
    - Table detection and recreation
    - Paragraph and heading detection
    - Multi-column layout handling
    """

    def __init__(self, pdf_path: str):
        self.pdf_path = pdf_path
        self.doc = fitz.open(pdf_path)
        self.word_doc = Document()

        # Set default styles
        self._setup_styles()

    def _setup_styles(self):
        """Configure default Word document styles"""
        from docx.shared import Twips
        from docx.enum.style import WD_STYLE_TYPE

        # Set default font
        style = self.word_doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Set paragraph spacing for Normal style (tighter spacing)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)  # Small gap between paragraphs
        style.paragraph_format.line_spacing = 1.15  # Slightly tighter line spacing

        # Set narrow margins for better layout matching
        sections = self.word_doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

    def _extract_text_blocks(self, page: fitz.Page) -> List[TextBlock]:
        """Extract all text blocks with formatting from a page"""
        blocks = []

        # Get text with detailed information
        text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

        for block_no, block in enumerate(text_dict.get("blocks", [])):
            if block.get("type") != 0:  # Skip non-text blocks
                continue

            for line_no, line in enumerate(block.get("lines", [])):
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if not text:
                        continue

                    # Extract formatting
                    font = span.get("font", "")
                    size = span.get("size", 11)
                    color = span.get("color", 0)
                    flags = span.get("flags", 0)
                    bbox = span.get("bbox", (0, 0, 0, 0))

                    # Parse color (integer to RGB)
                    r = ((color >> 16) & 0xFF) / 255.0
                    g = ((color >> 8) & 0xFF) / 255.0
                    b = (color & 0xFF) / 255.0

                    # Detect bold/italic from flags and font name
                    is_bold = bool(flags & 2**4) or "bold" in font.lower()
                    is_italic = bool(flags & 2**1) or "italic" in font.lower() or "oblique" in font.lower()

                    blocks.append(TextBlock(
                        text=span.get("text", ""),
                        x0=bbox[0],
                        y0=bbox[1],
                        x1=bbox[2],
                        y1=bbox[3],
                        font_name=font,
                        font_size=size,
                        color=(r, g, b),
                        is_bold=is_bold,
                        is_italic=is_italic,
                        block_no=block_no,
                        line_no=line_no
                    ))

        return blocks

    def _extract_images(self, page: fitz.Page) -> List[ImageBlock]:
        """Extract all images from a page"""
        images = []

        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]

            try:
                # Extract image
                base_image = self.doc.extract_image(xref)
                image_data = base_image["image"]
                ext = base_image["ext"]

                # Get image position
                img_rect = page.get_image_bbox(img)

                if img_rect:
                    images.append(ImageBlock(
                        image_data=image_data,
                        x0=img_rect.x0,
                        y0=img_rect.y0,
                        x1=img_rect.x1,
                        y1=img_rect.y1,
                        width=img_rect.width,
                        height=img_rect.height,
                        ext=ext
                    ))
            except Exception as e:
                print(f"Warning: Could not extract image {img_index}: {e}")
                continue

        return images

    def _detect_tables(self, page: fitz.Page) -> List[List[TableCell]]:
        """Detect and extract tables from a page with colors"""
        tables = []

        # Use PyMuPDF's table detection
        try:
            tab_finder = page.find_tables()

            # Get all drawings (rectangles with fill colors) for background detection
            drawings = page.get_drawings()

            # Get text dict for color extraction
            text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

            for table in tab_finder.tables:
                cells = []
                table_data = table.extract()

                # Get cell positions if available
                try:
                    cell_positions = table.cells  # List of cell rectangles
                except:
                    cell_positions = None

                for row_idx, row in enumerate(table_data):
                    for col_idx, cell_text in enumerate(row):
                        cell_x0, cell_y0, cell_x1, cell_y1 = 0, 0, 0, 0
                        bg_color = None
                        text_color = None
                        font_size = 11
                        is_bold = False

                        # Try to get cell position
                        if cell_positions:
                            try:
                                cell_idx = row_idx * len(row) + col_idx
                                if cell_idx < len(cell_positions):
                                    cell_rect = cell_positions[cell_idx]
                                    cell_x0, cell_y0, cell_x1, cell_y1 = cell_rect[:4]
                            except:
                                pass

                        # If we have cell position, look for background color
                        if cell_x0 != 0 or cell_y0 != 0:
                            # Find fill rectangles that overlap this cell
                            for drawing in drawings:
                                if drawing.get("fill") and drawing.get("rect"):
                                    d_rect = drawing["rect"]
                                    # Check if this drawing overlaps with the cell
                                    if (d_rect[0] <= cell_x0 + 5 and d_rect[2] >= cell_x1 - 5 and
                                        d_rect[1] <= cell_y0 + 5 and d_rect[3] >= cell_y1 - 5):
                                        fill_color = drawing.get("fill")
                                        if fill_color and fill_color != (1, 1, 1):  # Not white
                                            bg_color = fill_color
                                            break

                            # Find text color in this cell area
                            for block in text_dict.get("blocks", []):
                                if block.get("type") != 0:
                                    continue
                                for line in block.get("lines", []):
                                    for span in line.get("spans", []):
                                        span_bbox = span.get("bbox", (0, 0, 0, 0))
                                        # Check if span is within cell
                                        if (span_bbox[0] >= cell_x0 - 5 and span_bbox[2] <= cell_x1 + 5 and
                                            span_bbox[1] >= cell_y0 - 5 and span_bbox[3] <= cell_y1 + 5):
                                            # Get text color - ALWAYS preserve exact color
                                            color_int = span.get("color", 0)
                                            r = ((color_int >> 16) & 0xFF) / 255.0
                                            g = ((color_int >> 8) & 0xFF) / 255.0
                                            b = (color_int & 0xFF) / 255.0
                                            text_color = (r, g, b)  # Keep ALL colors including black

                                            # Get font info
                                            font_size = span.get("size", 11)
                                            flags = span.get("flags", 0)
                                            font_name = span.get("font", "")
                                            is_bold = bool(flags & 2**4) or "bold" in font_name.lower()
                                            break

                        cells.append(TableCell(
                            text=str(cell_text) if cell_text else "",
                            row=row_idx,
                            col=col_idx,
                            x0=cell_x0,
                            y0=cell_y0,
                            x1=cell_x1,
                            y1=cell_y1,
                            bg_color=bg_color,
                            text_color=text_color,
                            font_size=font_size,
                            is_bold=is_bold
                        ))

                if cells:
                    tables.append({
                        'cells': cells,
                        'rows': len(table_data),
                        'cols': len(table_data[0]) if table_data else 0,
                        'bbox': table.bbox
                    })
        except Exception as e:
            print(f"Warning: Table detection failed: {e}")

        return tables

    def _is_heading(self, block: TextBlock, avg_font_size: float) -> bool:
        """Determine if a text block is a heading"""
        # Heading indicators:
        # - Larger font size
        # - Bold
        # - Short text
        # - Starts with number/letter pattern

        is_larger = block.font_size > avg_font_size * 1.2
        is_short = len(block.text) < 100
        is_bold = block.is_bold

        # Check for heading patterns
        heading_pattern = re.match(r'^(\d+\.?\s*|[A-Z]\.\s*|Chapter\s+\d+|Section\s+\d+)', block.text)

        return (is_larger or is_bold) and is_short

    def _group_text_into_paragraphs(self, blocks: List[TextBlock]) -> List[dict]:
        """Group text blocks into paragraphs based on position, with spacing info"""
        if not blocks:
            return []

        # Sort by position (top to bottom, left to right)
        sorted_blocks = sorted(blocks, key=lambda b: (round(b.y0 / 10), b.x0))

        paragraphs = []
        current_para = [sorted_blocks[0]]
        prev_para_y1 = 0  # Track end of previous paragraph

        for i in range(1, len(sorted_blocks)):
            prev = sorted_blocks[i - 1]
            curr = sorted_blocks[i]

            # Check if same line (similar y position)
            same_line = abs(curr.y0 - prev.y0) < prev.font_size * 0.5

            # Check if continuation (small gap)
            small_gap = curr.y0 - prev.y1 < prev.font_size * 1.5

            # Check if same block
            same_block = curr.block_no == prev.block_no

            if same_line or (small_gap and same_block):
                current_para.append(curr)
            else:
                # Calculate spacing before this break
                para_y0 = min(b.y0 for b in current_para)
                para_y1 = max(b.y1 for b in current_para)
                space_before = para_y0 - prev_para_y1 if prev_para_y1 > 0 else 0

                paragraphs.append({
                    'blocks': current_para,
                    'y0': para_y0,
                    'y1': para_y1,
                    'space_before': space_before
                })
                prev_para_y1 = para_y1
                current_para = [curr]

        # Add last paragraph
        if current_para:
            para_y0 = min(b.y0 for b in current_para)
            para_y1 = max(b.y1 for b in current_para)
            space_before = para_y0 - prev_para_y1 if prev_para_y1 > 0 else 0
            paragraphs.append({
                'blocks': current_para,
                'y0': para_y0,
                'y1': para_y1,
                'space_before': space_before
            })

        return paragraphs

    def _add_formatted_text(self, paragraph, text: str, font_size: float,
                           is_bold: bool, is_italic: bool, color: Tuple[float, float, float]):
        """Add formatted text to a Word paragraph"""
        run = paragraph.add_run(text)
        run.font.size = Pt(font_size)
        run.font.bold = is_bold
        run.font.italic = is_italic

        # ALWAYS set color explicitly to preserve exact PDF formatting
        # This is important because heading styles may override with blue
        run.font.color.rgb = RGBColor(
            int(color[0] * 255),
            int(color[1] * 255),
            int(color[2] * 255)
        )

    def _add_image(self, image: ImageBlock, page_width: float):
        """Add an image to the Word document"""
        try:
            # Create image stream
            image_stream = io.BytesIO(image.image_data)

            # Calculate size (max width 6 inches to fit page)
            width_inches = min(image.width / 72, 6)  # Convert points to inches

            # Add to document
            self.word_doc.add_picture(image_stream, width=Inches(width_inches))

        except Exception as e:
            print(f"Warning: Could not add image: {e}")

    def _set_cell_shading(self, cell, color: Tuple[float, float, float]):
        """Set cell background color using XML"""
        try:
            # Convert RGB float (0-1) to hex
            r = int(color[0] * 255)
            g = int(color[1] * 255)
            b = int(color[2] * 255)
            hex_color = f"{r:02X}{g:02X}{b:02X}"

            # Get or create tcPr element
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            # Create shading element
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), hex_color)
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')

            # Remove existing shading if present
            existing_shd = tcPr.find(qn('w:shd'))
            if existing_shd is not None:
                tcPr.remove(existing_shd)

            tcPr.append(shading)
        except Exception as e:
            print(f"Warning: Could not set cell shading: {e}")

    def _add_table(self, table_data: dict):
        """Add a table to the Word document with colors"""
        try:
            rows = table_data['rows']
            cols = table_data['cols']
            cells = table_data['cells']

            if rows == 0 or cols == 0:
                return

            # Create table
            table = self.word_doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Fill cells with text, colors, and formatting
            for cell in cells:
                try:
                    table_cell = table.cell(cell.row, cell.col)

                    # Clear default paragraph and add formatted text
                    table_cell.text = ""  # Clear first
                    para = table_cell.paragraphs[0]

                    # Add text with formatting
                    if cell.text:
                        run = para.add_run(cell.text)

                        # Apply font size
                        if cell.font_size:
                            run.font.size = Pt(cell.font_size)

                        # Apply bold
                        if cell.is_bold:
                            run.font.bold = True

                        # ALWAYS apply text color to preserve exact formatting
                        if cell.text_color:
                            run.font.color.rgb = RGBColor(
                                int(cell.text_color[0] * 255),
                                int(cell.text_color[1] * 255),
                                int(cell.text_color[2] * 255)
                            )
                        else:
                            # Default to black if no color extracted
                            run.font.color.rgb = RGBColor(0, 0, 0)

                    # Apply background color
                    if cell.bg_color:
                        self._set_cell_shading(table_cell, cell.bg_color)

                except Exception as e:
                    print(f"Warning: Could not format cell ({cell.row}, {cell.col}): {e}")
                    continue

            # Add spacing after table
            self.word_doc.add_paragraph()

        except Exception as e:
            print(f"Warning: Could not add table: {e}")

    def convert(self, output_path: str = None) -> str:
        """
        Convert PDF to Word document.

        Args:
            output_path: Output .docx path. If None, uses input name with .docx extension.

        Returns:
            Path to the created Word document.
        """
        if output_path is None:
            output_path = str(Path(self.pdf_path).with_suffix('.docx'))

        # Process each page
        for page_num in range(len(self.doc)):
            page = self.doc[page_num]
            page_width = page.rect.width

            # Extract content
            text_blocks = self._extract_text_blocks(page)
            images = self._extract_images(page)
            tables = self._detect_tables(page)

            # Calculate average font size for heading detection
            if text_blocks:
                avg_font_size = sum(b.font_size for b in text_blocks) / len(text_blocks)
            else:
                avg_font_size = 11

            # Get table bounding boxes to exclude from text processing
            table_bboxes = [t['bbox'] for t in tables]

            # Filter out text that's inside tables
            def is_in_table(block: TextBlock) -> bool:
                for bbox in table_bboxes:
                    if (bbox[0] <= block.x0 <= bbox[2] and
                        bbox[1] <= block.y0 <= bbox[3]):
                        return True
                return False

            text_blocks = [b for b in text_blocks if not is_in_table(b)]

            # Group text into paragraphs
            paragraphs = self._group_text_into_paragraphs(text_blocks)

            # Combine all content with positions for ordering
            content_items = []

            # Add paragraphs (now using dict structure)
            for para_data in paragraphs:
                if para_data and para_data.get('blocks'):
                    content_items.append(('text', para_data['y0'], para_data))

            # Add images
            for img in images:
                content_items.append(('image', img.y0, img))

            # Add tables
            for table in tables:
                content_items.append(('table', table['bbox'][1], table))

            # Sort by vertical position
            content_items.sort(key=lambda x: x[1])

            # Add content to Word document
            for item_type, _, item in content_items:
                if item_type == 'text':
                    para_data = item
                    para_blocks = para_data['blocks']
                    space_before = para_data.get('space_before', 0)

                    # Create paragraph - DO NOT use Word's heading styles
                    # as they override text colors. Instead, preserve exact PDF formatting.
                    para = self.word_doc.add_paragraph()

                    # Set paragraph spacing based on PDF layout
                    # Convert PDF points to Word points (approximate)
                    if space_before > 20:  # Large gap - probably new section
                        para.paragraph_format.space_before = Pt(12)
                    elif space_before > 10:  # Medium gap
                        para.paragraph_format.space_before = Pt(6)
                    else:  # Small gap - same section
                        para.paragraph_format.space_before = Pt(0)

                    # Add text with formatting
                    for block in para_blocks:
                        self._add_formatted_text(
                            para,
                            block.text,
                            block.font_size,
                            block.is_bold,
                            block.is_italic,
                            block.color
                        )
                        # Add space between spans on same line
                        if block != para_blocks[-1]:
                            para.add_run(" ")

                elif item_type == 'image':
                    self._add_image(item, page_width)

                elif item_type == 'table':
                    self._add_table(item)

            # Add page break between pages (except last page)
            if page_num < len(self.doc) - 1:
                self.word_doc.add_page_break()

        # Save document
        self.word_doc.save(output_path)
        self.doc.close()

        return output_path

    def convert_to_bytes(self) -> bytes:
        """
        Convert PDF to Word and return as bytes.

        Returns:
            Word document as bytes
        """
        # Process the PDF
        for page_num in range(len(self.doc)):
            page = self.doc[page_num]
            page_width = page.rect.width

            text_blocks = self._extract_text_blocks(page)
            images = self._extract_images(page)
            tables = self._detect_tables(page)

            if text_blocks:
                avg_font_size = sum(b.font_size for b in text_blocks) / len(text_blocks)
            else:
                avg_font_size = 11

            table_bboxes = [t['bbox'] for t in tables]

            def is_in_table(block: TextBlock) -> bool:
                for bbox in table_bboxes:
                    if (bbox[0] <= block.x0 <= bbox[2] and
                        bbox[1] <= block.y0 <= bbox[3]):
                        return True
                return False

            text_blocks = [b for b in text_blocks if not is_in_table(b)]
            paragraphs = self._group_text_into_paragraphs(text_blocks)

            content_items = []

            # Add paragraphs (now using dict structure)
            for para_data in paragraphs:
                if para_data and para_data.get('blocks'):
                    content_items.append(('text', para_data['y0'], para_data))

            for img in images:
                content_items.append(('image', img.y0, img))

            for table in tables:
                content_items.append(('table', table['bbox'][1], table))

            content_items.sort(key=lambda x: x[1])

            for item_type, _, item in content_items:
                if item_type == 'text':
                    para_data = item
                    para_blocks = para_data['blocks']
                    space_before = para_data.get('space_before', 0)

                    # Create paragraph - DO NOT use Word's heading styles
                    # as they override text colors. Preserve exact PDF formatting.
                    para = self.word_doc.add_paragraph()

                    # Set paragraph spacing based on PDF layout
                    if space_before > 20:  # Large gap
                        para.paragraph_format.space_before = Pt(12)
                    elif space_before > 10:  # Medium gap
                        para.paragraph_format.space_before = Pt(6)
                    else:  # Small gap
                        para.paragraph_format.space_before = Pt(0)

                    for block in para_blocks:
                        self._add_formatted_text(
                            para,
                            block.text,
                            block.font_size,
                            block.is_bold,
                            block.is_italic,
                            block.color
                        )
                        if block != para_blocks[-1]:
                            para.add_run(" ")

                elif item_type == 'image':
                    self._add_image(item, page_width)

                elif item_type == 'table':
                    self._add_table(item)

            if page_num < len(self.doc) - 1:
                self.word_doc.add_page_break()

        # Save to bytes
        output = io.BytesIO()
        self.word_doc.save(output)
        output.seek(0)

        self.doc.close()

        return output.read()


def convert_pdf_to_word(input_path: str, output_path: str = None) -> str:
    """
    Convenience function to convert PDF to Word.

    Args:
        input_path: Path to input PDF file
        output_path: Path to output DOCX file (optional)

    Returns:
        Path to created Word document
    """
    converter = PDFToWordConverter(input_path)
    return converter.convert(output_path)


def convert_pdf_bytes_to_word_bytes(pdf_bytes: bytes) -> bytes:
    """
    Convert PDF bytes to Word document bytes.

    Args:
        pdf_bytes: PDF file as bytes

    Returns:
        Word document as bytes
    """
    # Create temp file-like object
    pdf_stream = io.BytesIO(pdf_bytes)
    doc = fitz.open(stream=pdf_stream, filetype="pdf")

    # Create converter with the document
    converter = PDFToWordConverter.__new__(PDFToWordConverter)
    converter.doc = doc
    converter.word_doc = Document()
    converter._setup_styles()

    return converter.convert_to_bytes()


# CLI interface
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python pdf_to_word.py <input.pdf> [output.docx]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    print(f"Converting {input_file}...")
    result = convert_pdf_to_word(input_file, output_file)
    print(f"Created: {result}")
